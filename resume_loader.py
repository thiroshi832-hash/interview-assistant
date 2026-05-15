"""
Resume loader. Accepts PDF / DOCX / TXT and returns plain text.
"""
from __future__ import annotations

from pathlib import Path


def load_resume(path: str | Path) -> str:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(p)

    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(p)
    if suffix == ".docx":
        return _load_docx(p)
    if suffix in (".txt", ".md"):
        return p.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"Unsupported resume format: {suffix} (use PDF, DOCX, or TXT)")


def _load_pdf(p: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(p))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n\n".join(parts).strip()


def _load_docx(p: Path) -> str:
    from docx import Document
    doc = Document(str(p))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # also pull tables, which resumes commonly use
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("  |  ".join(cells))
    return "\n".join(parts).strip()
